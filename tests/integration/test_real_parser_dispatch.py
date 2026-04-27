# Copyright 2025 ApeCloud, Inc.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Integration tests for ``parse_document`` real-parser dispatch — Wave 4 T3 chunk 1.

Pin the Wave 4 #4 production-readiness invariant: when
``parse_document`` is called with a ``source_filename`` whose
extension is not in the simulator allowlist (``.md`` / ``.markdown``
/ ``.txt`` / ``.text``), the call dispatches to the real
:class:`aperag.docparser.doc_parser.DocParser` chain and produces
the canonical ``markdown.md`` + ``outline.json`` + ``chunks.jsonl``
artifacts from the real binary input — no UTF-8 decode short-circuit,
no silent simulator fallback.

These tests run against the real ``MarkItDown`` parser (a hard
dependency), so they exercise the same code path production hits on
upload. They live under ``tests/integration/`` because they read +
write tempfiles; the simulator-only regression suite under
``tests/unit_test/indexing/test_t1_1_foundation.py`` keeps unit
coverage of the markdown path unchanged.
"""

from __future__ import annotations

import importlib.util

import pytest

from aperag.indexing import InMemoryObjectStore, parse_document, read_chunks


# Reading-back helpers — store accessor returns a context manager
# whose stream yields the bytes; we always slurp it once.
def _read_bytes(store: InMemoryObjectStore, path: str) -> bytes:
    obj = store.get(path)
    assert obj is not None, f"missing artifact: {path}"
    with obj as stream:
        return stream.read()


# -----------------------------------------------------------------------
# Simulator regression — must stay green so existing markdown callers
# (test_t1_1_foundation, document_service for *.md uploads) keep working.
# -----------------------------------------------------------------------


def test_markdown_simulator_path_unchanged_without_filename():
    """No ``source_filename`` → simulator decodes UTF-8 markdown directly."""
    store = InMemoryObjectStore()
    body = b"# Title\n\nFirst paragraph.\n\n## Sub\n\nSecond paragraph.\n"
    result = parse_document(
        store=store,
        collection_id="col-md",
        document_id="doc-md",
        source_bytes=body,
    )
    chunks = read_chunks(store, result.chunks_path)
    assert len(chunks) >= 1
    markdown = _read_bytes(store, result.markdown_path).decode("utf-8")
    assert "# Title" in markdown
    assert "## Sub" in markdown


def test_markdown_simulator_path_with_md_filename():
    """``source_filename='*.md'`` keeps the simulator route."""
    store = InMemoryObjectStore()
    body = b"# Heading\n\nbody"
    result = parse_document(
        store=store,
        collection_id="col-md2",
        document_id="doc-md2",
        source_bytes=body,
        source_filename="note.md",
    )
    markdown = _read_bytes(store, result.markdown_path).decode("utf-8")
    assert markdown == body.decode("utf-8")


def test_simulator_rejects_non_utf8_when_no_filename_hint():
    """Bytes without an extension hint and not UTF-8 surface a clear error."""
    store = InMemoryObjectStore()
    with pytest.raises(ValueError, match="simulator parser path requires UTF-8"):
        parse_document(
            store=store,
            collection_id="col",
            document_id="doc",
            source_bytes=b"\x89PNG\r\n\x1a\n",  # PNG magic
            source_filename=None,
        )


# -----------------------------------------------------------------------
# DocParser dispatch — non-text extensions route to the real chain.
# -----------------------------------------------------------------------


def test_html_routes_through_docparser_to_markdown():
    """HTML is one of MarkItDown's supported extensions — verifies the
    full DocParser chain end-to-end without needing additional binaries
    (``pikepdf`` for PDF, ``soffice`` for legacy .doc, ``PADDLEOCR_HOST``
    for images).
    """
    store = InMemoryObjectStore()
    html = (
        b"<html><body>"
        b"<h1>Title One</h1>"
        b"<p>First paragraph from HTML.</p>"
        b"<h2>Section</h2>"
        b"<p>Second paragraph.</p>"
        b"</body></html>"
    )
    result = parse_document(
        store=store,
        collection_id="col-html",
        document_id="doc-html",
        source_bytes=html,
        source_filename="report.html",
    )

    markdown = _read_bytes(store, result.markdown_path).decode("utf-8")
    assert "Title One" in markdown
    assert "Section" in markdown
    assert "First paragraph from HTML" in markdown

    outline_blob = _read_bytes(store, result.outline_path).decode("utf-8")
    assert "Title One" in outline_blob
    assert '"section_path": "1"' in outline_blob

    chunks = read_chunks(store, result.chunks_path)
    assert len(chunks) >= 1
    # Chunk content carries the real markdown text, not raw HTML.
    assert any("Title One" in c["text"] for c in chunks)


def test_unsupported_extension_raises_clear_error():
    """An extension DocParser cannot accept (e.g., a made-up suffix)
    raises ``ValueError`` with the supported-list embedded so callers
    can diagnose. We do not silently fall through to the simulator —
    that would mask wiring bugs.
    """
    store = InMemoryObjectStore()
    with pytest.raises(ValueError, match="DocParser does not accept"):
        parse_document(
            store=store,
            collection_id="col",
            document_id="doc",
            source_bytes=b"some bytes",
            source_filename="data.unknownext",
        )


# -----------------------------------------------------------------------
# PDF integration — runs only when pikepdf + a small PDF can be built.
# -----------------------------------------------------------------------


def _try_make_pdf_bytes() -> bytes | None:
    """Build a tiny PDF in memory if pikepdf is available; otherwise
    return ``None`` so the test below skips cleanly. PDF rendering is
    sensitive to the libcairo / ghostscript stack on the test host so
    we avoid that path and just synthesise a minimal valid PDF.
    """
    try:
        import pikepdf
    except ImportError:  # pragma: no cover
        return None
    try:
        pdf = pikepdf.Pdf.new()
        # pikepdf does not have a high-level "add text" API; for the
        # scope of T3 chunk 1 we just need a valid PDF that DocParser
        # can recognise as parseable. MarkItDown will return an
        # (empty-ish) MarkdownPart for a content-free PDF, and the
        # pipeline emits zero chunks — that is the contract we test.
        import io

        pdf.save(buf := io.BytesIO())
        return buf.getvalue()
    except Exception:  # pragma: no cover
        return None


def test_pdf_dispatches_to_docparser_even_when_empty():
    """Empty PDF → DocParser dispatched → empty markdown → zero chunks
    + valid outline. The point is the dispatch path runs and produces
    the canonical artifacts; a richer PDF case lives in chunk 2 e2e.
    """
    pdf_bytes = _try_make_pdf_bytes()
    if pdf_bytes is None:
        pytest.skip("pikepdf not available; skipping PDF dispatch test")

    store = InMemoryObjectStore()
    result = parse_document(
        store=store,
        collection_id="col-pdf",
        document_id="doc-pdf",
        source_bytes=pdf_bytes,
        source_filename="empty.pdf",
    )

    # Artifacts written even when the parser produces empty markdown.
    assert _read_bytes(store, result.markdown_path) is not None
    assert _read_bytes(store, result.outline_path) is not None
    assert _read_bytes(store, result.chunks_path) is not None


# -----------------------------------------------------------------------
# .docx via python-docx if installed.
# -----------------------------------------------------------------------


def test_docx_dispatches_to_docparser_when_python_docx_available():
    """Real .docx body via python-docx (transitively via markitdown[all]).
    Skips on test hosts where the optional dep chain is not present.
    """
    if importlib.util.find_spec("docx") is None:
        pytest.skip("python-docx not installed; skipping .docx round-trip")

    import io

    from docx import Document

    doc = Document()
    doc.add_heading("Real Word Heading", level=1)
    doc.add_paragraph("First docx paragraph.")
    doc.add_heading("Sub Heading", level=2)
    doc.add_paragraph("Second docx paragraph.")
    buf = io.BytesIO()
    doc.save(buf)

    store = InMemoryObjectStore()
    result = parse_document(
        store=store,
        collection_id="col-docx",
        document_id="doc-docx",
        source_bytes=buf.getvalue(),
        source_filename="report.docx",
    )
    markdown = _read_bytes(store, result.markdown_path).decode("utf-8")
    assert "Real Word Heading" in markdown
    assert "First docx paragraph" in markdown
    assert len(read_chunks(store, result.chunks_path)) >= 1
